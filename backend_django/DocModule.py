from docx import Document
from docx.shared import Inches,Pt,Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT,WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
import json
'''
json:
{
     "Subject":"bala"           #科目
     "Date":"bala"              #日期
     "Type":"bala:              #类型
     "Choice":[                   #选择题
     {
        "Question":"bala",
        "ChoiceA": "bala",
        "ChoiceB": "bala",
        "ChoiceC": "bala",
        "ChoiceD": "bala",
        "Answer":"bala"
     }

     ],

     "Completion":[                 #填空题
     {
        "Question":"bala",
        "Answer":"bala"
     }

     ],

     "Calculation":[                #计算题
     {
        "Question":"bala",
        "Answer":"bala"        
     }

     ],

     "Essay":[                      #问答题
     {
        "Question":"bala",
        "Answer":"bala"        
     }

     ]
}
'''
#添加分数框
def addPoint(Cell):
    Table=Cell.add_table(rows=2,cols=2)
    Table.style='Table Grid'
    
    Table.rows[0].height=Cm(0.82)
    Table.rows[1].height=Cm(0.82)
    Table.columns[0].width=Cm(2.12)
    Table.columns[1].width=Cm(1.56)
    Table.autofit=False

    Para=Table.cell(0,0).paragraphs[0]
    Para.add_run('本题分数')
    Para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    Table.cell(0,0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    Para=Table.cell(1,0).paragraphs[0]
    Para.add_run('得   分')
    Para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    Table.cell(1,0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

def addChoiceQ(Cell,Json):                                                  #选择题
   
    Para=Cell.add_paragraph()
    Run=Para.add_run('一.选择题 ')
    Run.font.name=u'黑体'
    Run.font.size=Pt(13)
    Run.element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')
    Run.bold=True
    Run=Para.add_run('(本大题共10题，每小题2分，共计20分)')
    addPoint(Cell)
    i=1
    for Text in Json['Choice']:
        Para=Cell.add_paragraph(str(i)+'.'+Text['Question'])
        Para=Cell.add_paragraph('A.'+Text['ChoiceA'])
        Para=Cell.add_paragraph('B.'+Text['ChoiceB'])
        Para=Cell.add_paragraph('C.'+Text['ChoiceC'])
        Para=Cell.add_paragraph('D.'+Text['ChoiceD'])
        i+=1
    

def addCompletionQ(Cell,Json):                                                #填空题
    
    Para=Cell.add_paragraph()
    Run=Para.add_run('二.填空题 ')
    Run.font.name=u'黑体'
    Run.font.size=Pt(13)
    Run.element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')
    Run.bold=True
    Run=Para.add_run('(本大题共10题，每空1分，共计10分)')
    addPoint(Cell)
    i=1
    for Text in Json['Completion']:
        Para=Cell.add_paragraph(str(i)+'.'+Text['Question'])
        i+=1
    

def addCalculationQ(Cell,Json):
    Para=Cell.add_paragraph()
    Run=Para.add_run('三.计算题 ')
    Run.font.name=u'黑体'
    Run.font.size=Pt(13)
    Run.element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')
    Run.bold=True
    Run=Para.add_run('(本大题共4题，每小题5分，共计20分)')
    i=1
    for Text in Json['Calculation']:
        addPoint(Cell)
        Para=Cell.add_paragraph(str(i)+'.'+Text['Question'])
        i+=1
    

def addEssayQ(Cell,Json):
    Para=Cell.add_paragraph()
    Run=Para.add_run('四.问答题 ')
    Run.font.name=u'黑体'
    Run.font.size=Pt(13)
    Run.element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')
    Run.bold=True
    Run=Para.add_run('(本大题共5题，每小题10分，共计50分)')
    i=1
    for Text in Json['Essay']:
        addPoint(Cell)
        Para=Cell.add_paragraph(str(i)+'.'+Text['Question'])
        i+=1
    

def addChoiceA(Cell,Json):
    Para=Cell.add_paragraph()
    Run=Para.add_run('一.选择题 ')
    Run.font.name=u'黑体'
    Run.font.size=Pt(13)
    Run.element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')
    Run.bold=True
    Run=Para.add_run('(本大题共10题，每小题2分，共计20分)')
    Para=Cell.add_paragraph()
    i=1
    for Text in Json['Choice']:
        Run=Para.add_run(str(i)+'.'+Text['Answer']+'.')
        i+=1


def addCompletionA(Cell,Json):
    Para=Cell.add_paragraph()
    Run=Para.add_run('二.填空题 ')
    Run.font.name=u'黑体'
    Run.font.size=Pt(13)
    Run.element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')
    Run.bold=True
    Run=Para.add_run('(本大题共10题，每空1分，共计10分)')
    Para=Cell.add_paragraph()
    i=1
    for Text in Json['Completion']:
        Run=Para.add_run(str(i)+'.'+Text['Answer'])
        i+=1


def addCalculationA(Cell,Json):
    Para=Cell.add_paragraph()
    Run=Para.add_run('三.计算题 ')
    Run.font.name=u'黑体'
    Run.font.size=Pt(13)
    Run.element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')
    Run.bold=True
    Run=Para.add_run('(本大题共4题，每小题5分，共计20分)')
    i=1
    for Text in Json['Calculation']:
        Para=Cell.add_paragraph(str(i)+'.'+Text['Answer'])
        i+=1


def addEssayA(Cell,Json):
    Para=Cell.add_paragraph()
    Run=Para.add_run('四.问答题 ')
    Run.font.name=u'黑体'
    Run.font.size=Pt(13)
    Run.element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')
    Run.bold=True
    Run=Para.add_run('(本大题共5题，每小题10分，共计50分)')
    i=1
    for Text in Json['Essay']:
        Para=Cell.add_paragraph(str(i)+'.'+Text['Answer'])
        i+=1
#调用自带的库
def genPaper(_Name='试卷.docx'):      #试卷生成
    Doc=Document('模板.docx')
    '''
    #标题 
    Title=Doc.paragraphs[0]   
    Title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    Title.paragraph_format.space_before=Pt(14)
    Run=Title.add_run('南 京 航 空 航 天 大 学')
    Run.font.name=u'楷体'
    Run.element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
    Run.font.size=Pt(22)
    Run.font.bold=True

    #页码
    Page=Doc.add_paragraph()
    Page.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    Run=Page.add_run('第1页 （共2页）')
    Run.font.size=Pt(10.5)
    Run.font.name=u'宋体'  
    Run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    Run.font.bold=False
    '''
  
    #整张卷子
    #卷头
    Table=Doc.tables[0]
    Json={"Date":"2016-7-8","Type":"A", "Subject":"数据结构","Choice":[  {"Question":"bala","ChoiceA": "bala","ChoiceB": "bala","ChoiceC": "bala", "ChoiceD": "bala" }],   "Completion":[ {"Question":"bala","Answer":"bala" } ],  "Calculation":[  { "Question":"bala",  "Answer":"bala"        }  ]}
    Para=Table.cell(0,0).paragraphs[0]
    Para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    N2C=["〇","一","二","三","四","五","六","七","八","九"]
    Date=Json['Date'].split('-')
    ToYear=int(Date[0])
    ToMonth=int(Date[1])
    ToDay=int(Date[2])
    if ToMonth<9:
        YearBefore=N2C[int((ToYear-1)/1000)%10]+N2C[int((ToYear-1)/100)%10]+N2C[int((ToYear-1)/10)%10]+N2C[int(ToYear-1)%10]
        YearAfter=N2C[int((ToYear)/1000)%10]+N2C[int((ToYear)/100)%10]+N2C[int((ToYear)/10)%10]+N2C[int(ToYear)%10]
        if ToMonth<3:
            Term=1
        else:
            Term=2
    else :
        YearBefore=N2C[int((ToYear)/1000)%10]+N2C[int((ToYear)/100)%10]+N2C[int((ToYear)/10)%10]+N2C[int(ToYear)%10]
        YearAfter=N2C[int((ToYear+1)/1000)%10]+N2C[int((ToYear+1)/100)%10]+N2C[int((ToYear+1)/10)%10]+N2C[int(ToYear+1)%10]
        Term=2

    Run=Para.add_run('%s ～ %s 学年  第 %d 学期'%(YearBefore,YearAfter,Term))         #todo
    Run.font.size=Pt(10.5) 
    Run.font.name=u'宋体'
    Run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    Run.bold=False

    Subject=Json['Subject']
    Run=Para.add_run('《%s》考试试题'%(Subject))
    Run.font.size=Pt(18) 
    Run.font.name=u'楷体'
    Run.element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
    Run.bold=True

    Para=Table.cell(0,0).paragraphs[1]   
    Type=Json['Type']
    Run=Para.add_run('考试日期：%d年%d月%d日     试卷类型：%s       试卷代号：'%(ToYear,ToMonth,ToDay,Type))
    Run.font.size=Pt(10.5) 
    Run.font.name=u'宋体'
    Run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    Run.bold=False

 
    #添加题目
    addChoiceQ(Table.cell(4,0),Json)
    addCompletionQ(Table.cell(4,0),Json)
    addCalculationQ(Table.cell(4,0),Json)
    addEssayQ(Table.cell(4,0),Json)
    #todo

    Doc.save(_Name)
def genAnswer(_Name='答案.docx'):     #答案生成
    Doc=Document('答案模板.docx')

    #卷头
    Table=Doc.tables[0]
    Json={"Date":"2016-7-8","Type":"A", "Subject":"数据结构","Choice":[  {"Question":"bala","ChoiceA": "bala","ChoiceB": "bala","ChoiceC": "bala", "ChoiceD": "bala" ,
        "Answer":"bala"}],   "Completion":[ {"Question":"bala","Answer":"bala" } ],  "Calculation":[  { "Question":"bala",  "Answer":"bala"        }  ],

     "Essay":[                   
     {
        "Question":"bala",
        "Answer":"bala"        
     }

     ]}
    N2C=["〇","一","二","三","四","五","六","七","八","九"]
    Date=Json['Date'].split('-')
    ToYear=int(Date[0])
    ToMonth=int(Date[1])
    ToDay=int(Date[2])
    if ToMonth<9:
        YearBefore=N2C[int((ToYear-1)/1000)%10]+N2C[int((ToYear-1)/100)%10]+N2C[int((ToYear-1)/10)%10]+N2C[int(ToYear-1)%10]
        YearAfter=N2C[int((ToYear)/1000)%10]+N2C[int((ToYear)/100)%10]+N2C[int((ToYear)/10)%10]+N2C[int(ToYear)%10]
        if ToMonth<3:
            Term=1
        else:
            Term=2
    else :
        YearBefore=N2C[int((ToYear)/1000)%10]+N2C[int((ToYear)/100)%10]+N2C[int((ToYear)/10)%10]+N2C[int(ToYear)%10]
        YearAfter=N2C[int((ToYear+1)/1000)%10]+N2C[int((ToYear+1)/100)%10]+N2C[int((ToYear+1)/10)%10]+N2C[int(ToYear+1)%10]
        Term=2

    Para=Table.cell(0,0).paragraphs[0]
    Run=Para.add_run('%s ～ %s 学年   第 %d 学期'%(YearBefore,YearAfter,Term))
    Run.font.size=Pt(13)
    Run.font.name=u'宋体'
    Run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    Run.bold=False

    
    Para=Table.cell(0,0).paragraphs[1]
    Run=Para.add_run('课程名称：')
    Run.font.size=Pt(13)
    Run.font.name=u'宋体'
    Run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    Run.bold=False
    Subject=Json['Subject']
    Run=Para.add_run('《%s》参考答案及评分标准'%(Subject))
    Run.font.size=Pt(18)
    Run.font.name=u'楷体'
    Run.element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
    Run.bold=True

    Type=Json['Type']
    Para=Table.cell(0,0).paragraphs[2]
    Run=Para.add_run('命题教师：                 试卷类型：%s       试卷代号：'%(Type))
    Run.font.size=Pt(13)
    Run.font.name=u'宋体'
    Run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    Run.bold=False
    addChoiceA(Table.cell(1,0),Json)
    addCompletionA(Table.cell(1,0),Json)
    addCalculationA(Table.cell(1,0),Json)
    addEssayA(Table.cell(1,0),Json)
    #todo
    Doc.save(_Name)
#genPaper()
genAnswer()

'''                                                 #大杀器 调用默认word软件
import win32com
from win32com.client import Dispatch, constants
import os,sys
W = win32com.client.Dispatch('Word.Application')
W.Visible = 0
W.DisplayAlerts = 0
Doc = W.Documents.Open( '%s\\sjmb.docx' %sys.path[0])
Doc.Close()
W.Quit()
'''